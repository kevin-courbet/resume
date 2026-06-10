from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clear_body(document: Document) -> None:
    body = document._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def update_header(document: Document) -> None:
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                if "DDZ." in run.text:
                    run.text = " KCB. – Expert Software, Data & Analytics Platforms"


def add_paragraph(
    document: Document, text: str = "", style: str | None = None, justify: bool = False
):
    paragraph = document.add_paragraph(style=style)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        paragraph.add_run(text)
    return paragraph


def add_first_heading(document: Document, text: str) -> None:
    paragraph = add_paragraph(document, style="Heading 1")
    run = paragraph.add_run()
    run.add_break()
    run.add_break()
    run.add_text(text)


def add_heading(document: Document, text: str) -> None:
    add_paragraph(document, text, "Heading 1")


def add_bullet(
    document: Document, text: str, bold_prefix: str | None = None, justify: bool = False
) -> None:
    paragraph = add_paragraph(document, style="List Bullet", justify=justify)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_number(document: Document, text: str, justify: bool = True) -> None:
    add_paragraph(document, text, "List Number", justify=justify)


def add_profile(document: Document, text: str) -> None:
    paragraph = add_paragraph(document, justify=True)
    paragraph.add_run("Profil : ").bold = True
    paragraph.add_run(text)


def experience(document: Document, title: str, profile: str, bullets: list[str]) -> None:
    add_bullet(document, title, bold_prefix=title, justify=True)
    add_profile(document, profile)
    for item in bullets:
        add_number(document, item)


def build_aerow_docx(template_path: Path, output_path: Path) -> None:
    document = Document(template_path)
    update_header(document)
    clear_body(document)

    add_first_heading(document, "PROFIL")
    add_paragraph(
        document,
        "Lead Software & Data Engineering avec plus de 8 ans d'expérience dans la conception, "
        "l'industrialisation et le pilotage de produits data-intensifs, depuis l'idéation "
        "jusqu'à la production.",
        justify=True,
    )
    add_paragraph(
        document,
        "Il intervient à l'intersection du développement logiciel, des plateformes data, de "
        "l'analytics embarqué et du leadership technique. Ses réalisations récentes incluent "
        "la construction, from scratch, d'une plateforme Databricks Lakehouse et d'un socle "
        "analytics client traitant plus de 20 millions d'enregistrements par jour pour Blue "
        "Systems (Bolloré Group), ainsi que la livraison de solutions d'analytics embarqué "
        "pour Amex GBT Consulting.",
        justify=True,
    )
    add_paragraph(
        document,
        "Ancien Engineering Manager chez Amex GBT Consulting, il encadre des équipes "
        "distribuées et structure des environnements de delivery fiables, observables et "
        "orientés performance. Il utilise les pratiques d'ingénierie assistée par IA et les "
        "workflows agentiques avec une approche pragmatique, centrée sur l'architecture, "
        "la modélisation de données, les tests et le jugement produit.",
        justify=True,
    )

    add_heading(document, "FORMATION")
    add_bullet(
        document, "Arts et Métiers ParisTech - Diplôme d'ingénieur", "Arts et Métiers ParisTech"
    )
    add_bullet(
        document,
        "Budapest University of Technology and Economics - Double diplôme en génie mécanique",
        "Budapest University of Technology and Economics",
    )

    add_heading(document, "COMPÉTENCES")
    paragraph = add_paragraph(document)
    paragraph.add_run("Langues :").bold = True
    paragraph.add_run(" Français (natif), Anglais (professionnel)")
    paragraph = add_paragraph(document)
    paragraph.add_run("Expertises principales :").bold = True
    for item in [
        "Conception et industrialisation de plateformes data : Databricks Lakehouse, Apache Spark/PySpark, Delta Lake, Delta Live Tables, dbt, SQL, ELT/ETL et modélisation de données",
        "Analytics & BI : analytics embarqué, Power BI, Tableau, self-service analytics, data visualization et reporting client-facing",
        "Développement logiciel : Golang, Python, TypeScript, JavaScript, React, Vue.js, C#, VBA, API REST et architectures cloud/serverless",
        "Architecture cloud et delivery : AWS, CI/CD, tests, observabilité, performance, optimisation des coûts et delivery agile",
        "Leadership technique : roadmaps, arbitrages d'architecture, management d'équipe, mentoring, collaboration distribuée et pilotage stakeholder",
        "Ingénierie agentique : développement assisté par LLM, context engineering, agent harnesses, automatisation de workflows et conception de boucles fiables",
    ]:
        add_bullet(document, item, justify=True)
    paragraph = add_paragraph(document)
    paragraph.add_run("Technologies & Outils :").bold = True
    for item in [
        "Data : Databricks, Spark/PySpark, Delta Lake, Delta Live Tables, dbt, SQL, Power BI, Tableau",
        "Software : Go, Python, TypeScript, JavaScript, React, Vue.js, C#, VBA, REST APIs",
        "Cloud & delivery : AWS, CI/CD, observabilité, testing, performance engineering, architecture cost-aware",
        "AI-assisted engineering : LLM workflows, agent orchestration, context management, automation tooling",
    ]:
        add_bullet(document, item)

    add_heading(document, "EXPÉRIENCES PROFESSIONNELLES")
    experience(
        document,
        "Mars 2023 - Juin 2026 - Blue Systems (Bolloré Group) / VersaDev",
        "Lead Developer / Data Platform Engineering",
        [
            "Construction from scratch d’une plateforme Data Analytics en production traitant plus de 20M d’enregistrements par jour sous contraintes fortes de coûts.",
            "Conception et implémentation d’un Databricks Lakehouse avec pipelines Spark/PySpark, transformations dbt, analytics Power BI et insights client-facing.",
            "Développement d’API Go pour les workflows d’ingestion et d’exploitation, intégrées à Databricks SQL Warehouse et à des services cloud/serverless.",
            "Mise en place de patterns cost-aware pour l’ingestion, la transformation, le serving analytics et les opérations afin de rendre la plateforme viable en production.",
            "Documentation technique de l’implémentation Blue Systems : Lakehouse, SQL Warehouses, Delta Live Tables, Unity Catalog et API Go pour service data smart-city.",
        ],
    )
    experience(
        document,
        "Juin 2023 - Août 2023 - Amex GBT Consulting / VersaDev",
        "Freelance Consultant / Front-end Delivery",
        [
            "Refonte du portail analytics consommateur d’Amex GBT Consulting afin de renforcer le self-service, l’interactivité et l’engagement utilisateur.",
            "Coordination avec les équipes produit, développement et infrastructure en cycles sprint, livraison de fonctionnalités Vue.js end-to-end et support de handover.",
        ],
    )
    experience(
        document,
        "Mars 2021 - Mars 2023 - Amex GBT Consulting, Global Solutions",
        "Engineering Manager",
        [
            "Management d’une équipe globale et entièrement remote de plus de 7 développeurs livrant analytics embarqué, pipelines analytiques et produits web.",
            "Mise en œuvre d’une des plus grandes plateformes européennes d’analytics embarqué dans l’univers consulting.",
            "Structuration des pratiques de product ownership et de delivery ; contribution aux roadmaps technologiques court et long terme avec les directions métier.",
            "Pilotage de la delivery technique multi-produits tout en restant 20 à 30% hands-on sur web full-stack, JavaScript, Python, C#, SQL et outils analytics.",
            "Mentoring des développeurs, promotion des pratiques d’ingénierie, facilitation inter-équipes et suppression des blocages de delivery.",
        ],
    )
    experience(
        document,
        "Mars 2020 - Février 2021 - Amex GBT Consulting, Global Solutions",
        "Lead Developer",
        [
            "Lead du développement de produits analytics, web, cloud, serverless et data visualization pour les équipes consulting et leurs clients.",
            "Conception de solutions techniques reliant besoins métier, pipelines data, API et expériences analytics utilisateur.",
            "Mise en place de pratiques de développement, workflows de delivery et patterns d’architecture pour des solutions analytics progressivement productisées.",
        ],
    )
    experience(
        document,
        "Février 2018 - Mars 2020 - Amex GBT Consulting",
        "Project Manager, Digital Strategy & Data",
        [
            "Pilotage d’initiatives de transformation digitale et data pour Consulting, en traduisant les besoins métier en produits analytiques, workflows et automatisations.",
            "Construction et amélioration d’outils internes, modèles analytiques et systèmes de reporting avec Excel/VBA, Python, SQL et technologies web.",
            "Collaboration avec les consultants business travel sur l’optimisation de processus, les négociations fournisseurs et l’aide à la décision data-driven.",
        ],
    )
    experience(
        document,
        "2016 - 2018 - Amex GBT Consulting",
        "Consultant, Business Travel & Analytics",
        [
            "Réalisation d’analyses de consulting en business travel, optimisation de processus, négociations fournisseurs et reporting de performance.",
            "Développement de modèles analytiques avancés et d’outillages Excel/VBA, établissant les fondations de l’évolution vers des rôles software et data engineering.",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

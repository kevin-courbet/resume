"""Build CV artifacts from structured source content."""

from __future__ import annotations

import argparse
import shutil
import subprocess
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader, select_autoescape

from cv_builder.aerow import build_aerow_docx

ACCENT = RGBColor(36, 79, 143)
INK = RGBColor(23, 32, 51)
MUTED = RGBColor(90, 101, 120)
STANDARD_BASE_NAME = "Kevin_Courbet_CV_2026_with_threadmill"
AEROW_BASE_NAME = "Kevin_Courbet_CV_AEROW_2026"
REQUIRED_TOP_LEVEL_KEYS = {
    "additional",
    "education",
    "email",
    "experience_page_1",
    "experience_page_2",
    "expertise",
    "github",
    "github_url",
    "headline",
    "linkedin",
    "linkedin_url",
    "location",
    "name",
    "phone",
    "phone_uri",
    "profile",
    "projects",
    "website",
    "website_url",
}


def load_resume_data(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as file:
        data = yaml.safe_load(file)
    if not isinstance(data, dict):
        msg = f"Expected mapping at {path}"
        raise TypeError(msg)
    missing_keys = sorted(REQUIRED_TOP_LEVEL_KEYS - data.keys())
    if missing_keys:
        msg = f"Missing required resume keys in {path}: {', '.join(missing_keys)}"
        raise KeyError(msg)
    return data


def render_html(data: dict[str, Any], template_dir: Path, output_path: Path, theme: str) -> None:
    env = Environment(
        loader=FileSystemLoader(str(template_dir)),
        autoescape=select_autoescape(enabled_extensions=("html", "j2")),
        trim_blocks=True,
        lstrip_blocks=True,
    )
    template = env.get_template("resume.html.j2")
    output_path.write_text(template.render(**data, theme=theme), encoding="utf-8")


def set_document_defaults(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(1.15)
        section.bottom_margin = Cm(1.05)
        section.left_margin = Cm(1.35)
        section.right_margin = Cm(1.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    normal.font.color.rgb = INK


def add_page_break(document: Document) -> None:
    document.add_section(WD_SECTION_START.NEW_PAGE)
    for section in document.sections:
        section.top_margin = Cm(1.15)
        section.bottom_margin = Cm(1.05)
        section.left_margin = Cm(1.35)
        section.right_margin = Cm(1.35)


def add_header(document: Document, data: dict[str, Any], compact: bool = False) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells

    name_paragraph = left.paragraphs[0]
    name_run = name_paragraph.add_run(data["name"])
    name_run.bold = True
    name_run.font.size = Pt(20 if compact else 26)
    name_run.font.color.rgb = INK

    headline = left.add_paragraph()
    headline_run = headline.add_run(
        "Professional Experience continued" if compact else data["headline"]
    )
    headline_run.bold = True
    headline_run.font.size = Pt(9 if compact else 10)
    headline_run.font.color.rgb = ACCENT

    contact = right.paragraphs[0]
    contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lines = (
        [f"{data['email']} | {data['phone']}", data["linkedin"]]
        if compact
        else [
            data["location"],
            f"{data['phone']} | {data['email']}",
            f"{data['linkedin']} | {data['website']}",
            data["github"],
        ]
    )
    for index, line in enumerate(lines):
        if index:
            contact.add_run().add_break()
        run = contact.add_run(line)
        run.font.size = Pt(7.5)
        run.font.color.rgb = MUTED


def add_section_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.space_before = Pt(8)
    paragraph.space_after = Pt(3)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(8.2)
    run.font.color.rgb = ACCENT


def add_bullets(document: Document, bullets: list[str], compact: bool = False) -> None:
    for bullet in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.space_after = Pt(1.1 if compact else 2)
        paragraph.paragraph_format.line_spacing = 1.02
        run = paragraph.add_run(bullet)
        run.font.size = Pt(8.2 if compact else 8.5)


def add_role_header(document: Document, title: str, context: str, dates: str) -> None:
    table = document.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    title_paragraph = left.paragraphs[0]
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = INK
    if context:
        context_paragraph = left.add_paragraph()
        context_run = context_paragraph.add_run(context)
        context_run.font.size = Pt(7.8)
        context_run.font.color.rgb = MUTED

    date_paragraph = right.paragraphs[0]
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_paragraph.add_run(dates)
    date_run.bold = True
    date_run.font.size = Pt(7.8)
    date_run.font.color.rgb = ACCENT


def add_client_card(document: Document, client: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.25)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(1)
    title = paragraph.add_run(f"{client['title']} - {client['subtitle']}")
    title.bold = True
    title.font.size = Pt(8.9)
    title.font.color.rgb = INK
    paragraph.add_run("    ")
    dates = paragraph.add_run(client["dates"])
    dates.bold = True
    dates.font.size = Pt(7.8)
    dates.font.color.rgb = ACCENT
    add_bullets(document, client["bullets"], compact=True)


def add_expertise_grid(document: Document, expertise: list[dict[str, str]]) -> None:
    table = document.add_table(rows=3, cols=2)
    table.autofit = True
    for index, item in enumerate(expertise):
        cell = table.cell(index // 2, index % 2)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(1)
        label = paragraph.add_run(f"{item['label']}: ")
        label.bold = True
        label.font.size = Pt(8)
        text = paragraph.add_run(item["text"])
        text.font.size = Pt(8)


def build_docx(data: dict[str, Any], output_path: Path) -> None:
    document = Document()
    set_document_defaults(document)

    add_header(document, data)

    add_section_title(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.05
    profile.add_run(data["profile"]).font.size = Pt(8.6)

    add_section_title(document, "Core Expertise")
    add_expertise_grid(document, data["expertise"])

    add_section_title(document, "Professional Experience")
    for role in data["experience_page_1"]:
        add_role_header(document, role["title"], role["context"], role["dates"])
        add_bullets(document, role.get("bullets", []), compact=True)
        for client in role.get("clients", []):
            add_client_card(document, client)
        add_bullets(document, role.get("trailing_bullets", []), compact=True)

    add_page_break(document)
    for role in data["experience_page_2"]:
        add_role_header(document, role["title"], role["context"], role["dates"])
        add_bullets(document, role["bullets"], compact=True)

    add_section_title(document, "Selected Projects")
    for project in data["projects"]:
        add_role_header(document, project["title"], "", project["dates"])
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(project["text"]).font.size = Pt(8.3)

    add_section_title(document, "Education")
    for edu in data["education"]:
        paragraph = document.add_paragraph()
        title_run = paragraph.add_run(f"{edu['title']} - ")
        title_run.bold = True
        title_run.font.size = Pt(8.5)
        paragraph.add_run(edu["meta"]).font.size = Pt(8.5)

    add_section_title(document, "Additional")
    paragraph = document.add_paragraph()
    paragraph.add_run(data["additional"]).font.size = Pt(8.5)

    document.save(output_path)


def render_html_pdf(html_path: Path, output_path: Path) -> None:
    from weasyprint import HTML

    HTML(filename=str(html_path)).write_pdf(str(output_path))


def render_docx_pdf(docx_path: Path, output_path: Path) -> None:
    office_binary = shutil.which("soffice") or shutil.which("libreoffice")
    if office_binary is None:
        msg = "LibreOffice not found. Install LibreOffice or run with --skip-docx-pdf."
        raise FileNotFoundError(msg)

    output_dir = output_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            office_binary,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        check=True,
    )
    generated = output_dir / f"{docx_path.stem}.pdf"
    if generated != output_path:
        generated.replace(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build CV artifacts.")
    parser.add_argument("--data", type=Path, default=Path("data/resume.yml"))
    parser.add_argument("--template-dir", type=Path, default=Path("templates"))
    parser.add_argument("--out-dir", type=Path, default=Path("dist"))
    parser.add_argument("--base-name", default=STANDARD_BASE_NAME)
    parser.add_argument("--target", choices=("standard", "aerow", "all"), default="all")
    parser.add_argument("--theme", choices=("light", "dark", "all"), default="all")
    parser.add_argument("--skip-docx-pdf", action="store_true")
    parser.add_argument("--skip-html-pdf", action="store_true")
    return parser.parse_args()


def build_standard_artifacts(args: argparse.Namespace) -> list[Path]:
    docx_path = args.out_dir / f"{args.base_name}.docx"
    docx_pdf_path = args.out_dir / f"{args.base_name}_docx.pdf"
    themes = ("light", "dark") if args.theme == "all" else (args.theme,)
    generated_paths = [docx_path]

    data = load_resume_data(args.data)
    build_docx(data, docx_path)

    if not args.skip_docx_pdf:
        render_docx_pdf(docx_path, docx_pdf_path)
        generated_paths.append(docx_pdf_path)

    for theme in themes:
        html_path = args.out_dir / f"{args.base_name}_{theme}.html"
        html_pdf_path = args.out_dir / f"{args.base_name}_{theme}_html.pdf"
        render_html(data, args.template_dir, html_path, theme)
        generated_paths.append(html_path)
        if not args.skip_html_pdf:
            render_html_pdf(html_path, html_pdf_path)
            generated_paths.append(html_pdf_path)

    return generated_paths


def build_aerow_artifacts(args: argparse.Namespace) -> list[Path]:
    docx_path = args.out_dir / f"{AEROW_BASE_NAME}.docx"
    pdf_path = args.out_dir / f"{AEROW_BASE_NAME}.pdf"
    template_path = args.template_dir / "aerow-template.docx"
    generated_paths = [docx_path]

    build_aerow_docx(template_path, docx_path)
    if not args.skip_docx_pdf:
        render_docx_pdf(docx_path, pdf_path)
        generated_paths.append(pdf_path)

    return generated_paths


def main() -> None:
    args = parse_args()
    args.out_dir.mkdir(parents=True, exist_ok=True)

    generated_paths: list[Path] = []
    if args.target in {"standard", "all"}:
        generated_paths.extend(build_standard_artifacts(args))
    if args.target in {"aerow", "all"}:
        generated_paths.extend(build_aerow_artifacts(args))

    print("Generated:")
    for path in generated_paths:
        print(f"- {path}")


if __name__ == "__main__":
    main()
